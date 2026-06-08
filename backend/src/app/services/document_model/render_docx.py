from __future__ import annotations

import io
from typing import Any

from app.core.errors import ValidationAppError

Document = None
WD_SECTION = None
WD_ALIGN_PARAGRAPH = None
OxmlElement = None
qn = None
Inches = None
Pt = None
RGBColor = None


def render_docx(model: dict[str, Any]) -> bytes:
    global Document, WD_SECTION, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor
    try:
        from docx import Document as _Document
        from docx.enum.section import WD_SECTION as _WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor

        Document = _Document
        WD_SECTION = _WD_SECTION
        WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
        OxmlElement = _OxmlElement
        qn = _qn
        Inches = _Inches
        Pt = _Pt
        RGBColor = _RGBColor
    except Exception as exc:  # pragma: no cover
        raise ValidationAppError("python-docx is required to generate DOCX artifacts") from exc

    document = Document()
    _configure_document(document, qn, Inches)
    _configure_styles(document, qn, Pt, RGBColor)
    _header_footer(document, model, OxmlElement, qn)
    _cover(document, model, WD_ALIGN_PARAGRAPH)
    if (model.get("toc") or {}).get("enabled", True):
        _toc(document)
    for section in model.get("sections", []):
        _render_section(document, section)
    if model.get("appendix"):
        document.add_page_break()
        document.add_heading("附录", level=1)
        for section in model["appendix"]:
            _render_section(document, section)
    if model.get("signatures"):
        _signatures(document, model["signatures"])
    document.add_section(WD_SECTION.CONTINUOUS)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_document(document, qn, inches) -> None:
    for section in document.sections:
        section.top_margin = inches(0.72)
        section.bottom_margin = inches(0.68)
        section.left_margin = inches(0.82)
        section.right_margin = inches(0.82)
    settings = document.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is not None:
        compat.set(qn("w:compatSetting"), "true")


def _configure_styles(document, qn, pt_factory, color_factory) -> None:
    font_name = "Microsoft YaHei"
    normal = document.styles["Normal"]
    _style_font(normal, qn, font_name, pt_factory(10.5), color_factory(55, 65, 81))
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = pt_factory(6)
    for style_name, size in (("Title", 28), ("Heading 1", 19), ("Heading 2", 15.5), ("Heading 3", 13)):
        style = document.styles[style_name]
        _style_font(style, qn, font_name, pt_factory(size), color_factory(17, 24, 39))
        style.paragraph_format.space_before = pt_factory(10)
        style.paragraph_format.space_after = pt_factory(6)
    for style_name in ("List Bullet", "List Number", "Intense Quote"):
        _style_font(document.styles[style_name], qn, font_name, pt_factory(10.5), color_factory(55, 65, 81))

    try:
        table_style = document.styles["Table Grid"]
        _style_font(table_style, qn, font_name, pt_factory(10), color_factory(55, 65, 81))
    except Exception:
        pass


def _style_font(style, qn, name: str, size, color) -> None:
    style.font.name = name
    style.font.size = size
    style.font.color.rgb = color
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _header_footer(document, model: dict[str, Any], element_factory, qn) -> None:
    title = str(model.get("title") or "AgentHub 文档")
    header = str((model.get("template_spec") or {}).get("header") or f"AgentHub · {title}")
    footer_text = str((model.get("template_spec") or {}).get("footer") or "页码")
    for section in document.sections:
        header_para = section.header.paragraphs[0]
        header_para.text = f"{header} · {title}"
        header_para.alignment = 1
        footer = section.footer.paragraphs[0]
        footer.alignment = 1
        footer.text = f"{footer_text} · 第 "
        run = footer.add_run()
        _append_page_field(run, element_factory, qn)
        footer.add_run(" 页")


def _append_page_field(run, element_factory, qn) -> None:
    fld_begin = element_factory("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = element_factory("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = element_factory("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _cover(document, model: dict[str, Any], align) -> None:
    cover = model.get("cover") or {}
    title = document.add_paragraph()
    title.alignment = align.CENTER
    title.style = document.styles["Title"]
    title.add_run(str(cover.get("title") or model.get("title") or "AgentHub 文档")).bold = True
    subtitle = str(cover.get("subtitle") or model.get("subtitle") or "")
    if subtitle:
        para = document.add_paragraph(subtitle)
        para.alignment = align.CENTER
        para.paragraph_format.space_after = Pt(10)
    kicker = document.add_paragraph()
    kicker.alignment = align.CENTER
    kicker.add_run("AGENTHUB DOCUMENT").bold = True
    meta_lines = [
        str(cover.get("issuer") or "AgentHub"),
        str(cover.get("confidentiality") or ""),
        f"{cover.get('date_label') or '生成日期'}：{cover.get('date') or ''}",
    ]
    for line in [item for item in meta_lines if item]:
        para = document.add_paragraph(line)
        para.alignment = align.CENTER
        para.paragraph_format.space_after = Pt(4)
    line_para = document.add_paragraph("结构化生成 · 可预览 · 可导出 · 可归档")
    line_para.alignment = align.CENTER
    document.add_page_break()


def _toc(document) -> None:
    document.add_heading("目录", level=1)
    hint = document.add_paragraph("请在 Word/WPS 中右键更新目录，以生成精确页码。")
    hint.style = document.styles["Intense Quote"]
    document.add_page_break()


def _render_section(document, section: dict[str, Any]) -> None:
    title = str(section.get("title") or "")
    if title:
        document.add_heading(title, level=int(section.get("level") or 1))
    for block in section.get("blocks", []):
        _render_block(document, block)


def _render_block(document, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    if block_type == "heading":
        document.add_heading(str(block.get("text") or ""), level=int(block.get("level") or 2))
    elif block_type == "list":
        _render_list(document, block)
    elif block_type == "table":
        _render_table(document, block)
    elif block_type == "callout":
        para = document.add_paragraph(style="Intense Quote")
        para.add_run(f"{block.get('title') or '提示'}：").bold = True
        para.add_run(str(block.get("text") or ""))
    elif block_type == "quote":
        document.add_paragraph(str(block.get("text") or ""), style="Intense Quote")
    elif block_type == "image":
        document.add_paragraph(f"[图片占位：{block.get('alt') or block.get('src') or '未命名图片'}]")
    elif block_type == "divider":
        divider = document.add_paragraph("—" * 36)
        divider.alignment = 1
    elif block_type == "page_break":
        document.add_page_break()
    elif block_type == "signatures":
        _signatures(document, block.get("items") or [])
    else:
        document.add_paragraph(str(block.get("text") or ""))


def _render_list(document, block: dict[str, Any]) -> None:
    style = "List Number" if block.get("ordered") else "List Bullet"
    for item in block.get("items", []):
        document.add_paragraph(str(item), style=style)


def _render_table(document, block: dict[str, Any]) -> None:
    headers = [str(item) for item in block.get("headers", [])]
    rows = [[str(cell) for cell in row] for row in block.get("rows", [])]
    if not headers and not rows:
        return
    column_count = max(len(headers), *(len(row) for row in rows), 1)
    table = document.add_table(rows=1 if headers else 0, cols=column_count)
    table.style = "Table Grid"
    if headers:
        for index, value in enumerate(headers):
            table.rows[0].cells[index].text = value
            _format_table_cell(table.rows[0].cells[index], header=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values[:column_count]):
            cells[index].text = value
            _format_table_cell(cells[index], header=False)
    _style_table(table)


def _signatures(document, items: list[str]) -> None:
    names = [str(item) for item in items if str(item).strip()] or ["负责人", "确认人"]
    document.add_heading("签字确认", level=1)
    table = document.add_table(rows=len(names) + 1, cols=3)
    table.style = "Table Grid"
    for index, header in enumerate(["角色", "签字", "日期"]):
        table.rows[0].cells[index].text = header
        _format_table_cell(table.rows[0].cells[index], header=True)
    for row_index, name in enumerate(names, start=1):
        table.rows[row_index].cells[0].text = name
        table.rows[row_index].cells[1].text = ""
        table.rows[row_index].cells[2].text = ""
        for cell in table.rows[row_index].cells:
            _format_table_cell(cell, header=False)


def _style_table(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                _set_cell_background(cell, "D9ECFF")
            else:
                _set_cell_background(cell, "F8FAFC" if row_index % 2 == 0 else "FFFFFF")


def _format_table_cell(cell, *, header: bool) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            run.bold = header
            if header:
                run.font.color.rgb = RGBColor(15, 23, 42)
    _set_cell_margins(cell, top=120, start=120, bottom=120, end=120)


def _set_cell_background(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
