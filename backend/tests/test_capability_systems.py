from __future__ import annotations

import io
import json
import threading
import zipfile
from contextlib import contextmanager
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from pypdf import PdfReader
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.core.errors import ValidationAppError
from db.base import Base
from db.models import (
    Agent,
    Artifact,
    ArtifactVersion,
    Conversation,
    FileAsset,
    McpServer,
    McpToolInvocation,
    Skill,
    ToolDefinition,
    ToolInvocation,
    User,
)
from app.services.agents.tool_loop import build_tools_for_agent, execute_tool_by_name
from app.services.tools.builtins.artifact.export import default_export_format, export_artifact
from app.services.artifacts import update_artifact_files
from app.services.demo_cleanup import cleanup_acceptance_residue
from app.services.mcp.discovery import discover_server_tools, import_server_manifest, probe_server, probe_server_async
from app.services.mcp.invocation import invoke_mcp_tool_recorded
from app.services.mcp.schema import validate_mcp_arguments
from app.services.mcp.transports.common import tool_allowed
from app.services.skills.context import activated_skill_context
from app.services.skills.package import parse_skill_package
from app.services.document_model import normalize_document_model, parse_markdown_blocks
from app.services.tools.executor import invoke_tool
from app.services.tools.catalog import sync_builtin_tool_definitions


@contextmanager
def _mcp_tools_list_server() -> Any:
    class Handler(BaseHTTPRequestHandler):
        def do_POST(self) -> None:  # noqa: N802 - stdlib handler API
            length = int(self.headers.get("Content-Length", "0"))
            body = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            if body.get("method") != "tools/list":
                self.send_response(404)
                self.end_headers()
                return
            data = json.dumps(
                {
                    "jsonrpc": "2.0",
                    "id": body.get("id"),
                    "result": {
                        "tools": [
                            {
                                "name": "echo.ping",
                                "description": "Echo input",
                                "inputSchema": {
                                    "type": "object",
                                    "properties": {"input": {"type": "string"}},
                                    "required": ["input"],
                                },
                            }
                        ]
                    },
                }
            ).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        def log_message(self, *_args: Any) -> None:
            return

    server = ThreadingHTTPServer(("127.0.0.1", 0), Handler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        yield f"http://127.0.0.1:{server.server_port}"
    finally:
        server.shutdown()
        thread.join(timeout=2)
        server.server_close()


def test_tool_invocation_records_builtin_run() -> None:
    db = _memory_session()
    user = _user()
    db.add(user)
    db.commit()

    result = invoke_tool(db, user, "api.test", {"path": "/api/v1/health"})
    invocation = db.scalar(select(ToolInvocation).where(ToolInvocation.id == result["invocation_id"]))

    assert result["result"]["status"] == "succeeded"
    assert invocation is not None
    assert invocation.tool_name == "api.test"
    assert invocation.status == "succeeded"


def test_artifact_create_pdf_exports_real_chinese_pdf() -> None:
    db = _memory_session()
    user = _user()
    conversation = Conversation(creator_id=user.id, chat_type="single", title="PDF Conversation")
    db.add_all([user, conversation])
    db.commit()

    result = invoke_tool(
        db,
        user,
        "artifact.create_pdf",
        {
            "conversation_id": conversation.id,
            "title": "中文项目方案",
            "body": "# 目标\n这是中文段落。\n- 第一项\n- 第二项\n---\n第二页内容",
        },
    )
    artifact = db.get(Artifact, result["result"]["artifact"]["id"])
    exported = export_artifact(artifact, "pdf")
    reader = PdfReader(io.BytesIO(exported.content))

    assert result["result"]["capability_level"] == "real"
    assert exported.media_type == "application/pdf"
    assert exported.content.startswith(b"%PDF")
    assert len(reader.pages) >= 2
    assert b"/FontName" in exported.content
    assert any(token in exported.content for token in (b"NotoSans", b"SimHei", b"STSong"))


@pytest.mark.parametrize(
    ("tool_name", "fmt"),
    [
        ("artifact.create_pdf", "pdf"),
        ("artifact.create_docx", "docx"),
        ("artifact.create_xlsx", "xlsx"),
        ("artifact.create_pptx", "pptx"),
    ],
)
def test_office_artifacts_persist_real_files_and_versions(tool_name: str, fmt: str) -> None:
    db = _memory_session()
    user = _user()
    conversation = Conversation(creator_id=user.id, chat_type="single", title="Office Files")
    db.add_all([user, conversation])
    db.commit()

    result = invoke_tool(
        db,
        user,
        tool_name,
        {"conversation_id": conversation.id, "title": f"真实{fmt}文件", "body": "第一段\n第二段"},
    )
    artifact = db.get(Artifact, result["result"]["artifact_id"])
    source_file = artifact.content["source_file"]
    exported = export_artifact(artifact, fmt)
    default_exported = export_artifact(artifact)
    asset = db.get(FileAsset, source_file["file_asset_id"])

    assert artifact.content["format"] == fmt
    assert default_export_format(artifact) == fmt
    assert artifact.content["preview_html"]
    assert "AgentHub Artifact" not in artifact.content["preview_html"]
    assert artifact.content["source_text"]
    assert artifact.content["content_model"]["kind"] in {"document", "spreadsheet", "slides"}
    assert source_file == artifact.content["export_file"]
    assert Path(source_file["storage_path"]).read_bytes() == exported.content
    assert default_exported.content == exported.content
    assert default_exported.filename.endswith(f".{fmt}")
    assert asset is not None
    assert asset.artifact_id == artifact.id
    assert asset.purpose == "artifact_source"
    assert asset.size == len(exported.content)

    update_artifact_files(db, artifact.id, {"index.html": "<main><h1>第二版</h1></main>"}, "保存第二版")
    db.refresh(artifact)
    version = db.scalar(
        select(ArtifactVersion).where(
            ArtifactVersion.artifact_id == artifact.id,
            ArtifactVersion.version == artifact.current_version,
        )
    )

    assert artifact.current_version == 2
    assert artifact.content["source_file"]["version"] == 2
    assert Path(artifact.content["source_file"]["storage_path"]).exists()
    assert version is not None
    assert version.checksum == artifact.content["source_file"]["checksum"]


def test_document_model_drives_docx_and_pdf_renderers() -> None:
    db = _memory_session()
    user = _user()
    conversation = Conversation(creator_id=user.id, chat_type="single", title="Document Model")
    db.add_all([user, conversation])
    db.commit()
    content_model = _document_model()

    docx_result = invoke_tool(
        db,
        user,
        "artifact.create_docx",
        {
            "conversation_id": conversation.id,
            "title": "结构化项目方案",
            "template": "proposal",
            "content_model": content_model,
        },
    )
    pdf_result = invoke_tool(
        db,
        user,
        "artifact.create_pdf",
        {
            "conversation_id": conversation.id,
            "title": "结构化项目方案",
            "template": "proposal",
            "content_model": content_model,
        },
    )

    docx_artifact = db.get(Artifact, docx_result["result"]["artifact_id"])
    pdf_artifact = db.get(Artifact, pdf_result["result"]["artifact_id"])
    docx_export = export_artifact(docx_artifact)
    pdf_export = export_artifact(pdf_artifact)
    document = Document(io.BytesIO(docx_export.content))
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf_export.content)).pages)

    assert docx_artifact.content["content_model"]["template"] == "proposal"
    assert "agenthub-word-preview" in docx_artifact.content["preview_html"]
    assert "callout" in docx_artifact.content["preview_html"]
    assert any(paragraph.text == "实施范围" for paragraph in document.paragraphs)
    assert document.tables and document.tables[0].cell(0, 0).text == "阶段"
    assert "结构化项目方案" in pdf_text
    assert "实施范围" in pdf_text
    assert "阶段" in pdf_text


def test_document_templates_have_distinct_default_sections() -> None:
    report = normalize_document_model(None, title="报告模板", source_text="", template="report")
    proposal = normalize_document_model(None, title="方案模板", source_text="", template="proposal")
    lab_report = normalize_document_model(None, title="实验模板", source_text="", template="lab_report")

    report_titles = [section["title"] for section in report["sections"]]
    proposal_titles = [section["title"] for section in proposal["sections"]]
    lab_titles = [section["title"] for section in lab_report["sections"]]

    assert "分析内容" in report_titles
    assert "实施计划" in proposal_titles
    assert "实验目的" in lab_titles
    assert report_titles != proposal_titles
    assert proposal_titles != lab_titles


def test_markdown_input_becomes_structured_document_blocks() -> None:
    markdown = """# 背景
第一行说明
第二行继续

- 事项 A
- 事项 B

| 指标 | 结果 |
| --- | --- |
| 通过率 | 100% |

> [!WARNING] 风险提示
> 需要持续观察。

> 这是一段引用

![架构图](architecture.png)

---

[page_break]
"""
    blocks = parse_markdown_blocks(markdown)
    model = normalize_document_model(None, title="Markdown 文档", source_text=markdown, template="report")
    flattened = [block for section in model["sections"] for block in section["blocks"]]

    assert any(block["type"] == "heading" and block["text"] == "背景" for block in blocks)
    assert any(block["type"] == "paragraph" and "第一行说明 第二行继续" in block["text"] for block in flattened)
    assert any(block["type"] == "list" and block["items"] == ["事项 A", "事项 B"] for block in flattened)
    assert any(block["type"] == "table" and block["headers"] == ["指标", "结果"] for block in flattened)
    assert any(block["type"] == "callout" and block["variant"] == "warning" for block in flattened)
    assert any(block["type"] == "quote" for block in flattened)
    assert any(block["type"] == "image" and block["src"] == "architecture.png" for block in flattened)
    assert any(block["type"] == "divider" for block in flattened)
    assert any(block["type"] == "page_break" for block in flattened)


@pytest.mark.parametrize(
    ("tool_name", "fmt", "preview_marker"),
    [
        ("artifact.create_docx", "docx", "agenthub-word-preview"),
        ("artifact.create_xlsx", "xlsx", "agenthub-sheet-preview"),
        ("artifact.create_pptx", "pptx", "agenthub-slides-preview"),
        ("artifact.create_pdf", "pdf", "agenthub-word-preview"),
    ],
)
def test_artifact_exports_cover_zip_and_text_formats(tool_name: str, fmt: str, preview_marker: str) -> None:
    db = _memory_session()
    user = _user()
    conversation = Conversation(creator_id=user.id, chat_type="single", title="Export Formats")
    db.add_all([user, conversation])
    db.commit()

    result = invoke_tool(
        db,
        user,
        tool_name,
        {"conversation_id": conversation.id, "title": "导出测试", "body": "目标\n计划\n验收"},
    )
    artifact = db.get(Artifact, result["result"]["artifact_id"])
    zip_export = export_artifact(artifact, "zip")
    html_export = export_artifact(artifact, "html")
    md_export = export_artifact(artifact, "markdown")
    json_export = export_artifact(artifact, "json")

    assert preview_marker in artifact.content["preview_html"]
    if fmt in {"docx", "pdf"}:
        assert "HTML preview" in artifact.content["preview_html"]
    assert "\ufffd" not in artifact.content["preview_html"]
    assert html_export.media_type.startswith("text/html")
    assert md_export.media_type.startswith("text/markdown")
    assert json_export.media_type.startswith("application/json")
    with zipfile.ZipFile(io.BytesIO(zip_export.content)) as archive:
        names = set(archive.namelist())
        source_name = f"source/{artifact.content['source_file']['filename']}"
        assert {"metadata.json", "README.md", "preview.html", source_name}.issubset(names)
        assert archive.read(source_name) == Path(artifact.content["source_file"]["storage_path"]).read_bytes()


def test_html_artifact_preview_and_export_use_real_html_file() -> None:
    db = _memory_session()
    user = _user()
    conversation = Conversation(creator_id=user.id, chat_type="single", title="HTML Files")
    db.add_all([user, conversation])
    db.commit()
    html = "<!doctype html><html><body><main><h1>HTML 预览</h1></main></body></html>"

    result = invoke_tool(
        db,
        user,
        "artifact.create_html",
        {"conversation_id": conversation.id, "title": "HTML 预览", "html": html},
    )
    artifact = db.get(Artifact, result["result"]["artifact_id"])
    exported = export_artifact(artifact, "html")

    assert artifact.content["preview_html"] == html
    assert artifact.content["source_file"]["format"] == "html"
    assert exported.content.decode("utf-8") == html
    assert Path(artifact.content["source_file"]["storage_path"]).read_text(encoding="utf-8") == html


def test_sandbox_run_timeout_and_denied_command_are_recorded() -> None:
    db = _memory_session()
    user = _user()
    db.add(user)
    db.commit()

    timeout_result = invoke_tool(
        db,
        user,
        "sandbox.run",
        {"command": 'python -c "while True: pass"', "timeout": 1},
    )

    assert timeout_result["result"]["status"] == "timeout"
    assert timeout_result["result"]["exit_code"] == -1
    assert timeout_result["result"]["capability_level"] == "real"

    with pytest.raises(ValidationAppError):
        invoke_tool(db, user, "sandbox.run", {"command": "rm -rf ."})
    failed = db.scalars(select(ToolInvocation).where(ToolInvocation.tool_name == "sandbox.run")).all()
    assert failed[-1].status == "failed"


def test_test_run_executes_allowed_command() -> None:
    db = _memory_session()
    user = _user()
    db.add(user)
    db.commit()

    result = invoke_tool(db, user, "test.run", {"command": "pytest --version", "timeout": 10})

    assert result["result"]["status"] == "succeeded"
    assert result["result"]["exit_code"] == 0
    assert "pytest" in result["result"]["stdout"].lower()


def test_api_test_asserts_status_code() -> None:
    db = _memory_session()
    user = _user()
    db.add(user)
    db.commit()

    result = invoke_tool(
        db,
        user,
        "api.test",
        {"method": "GET", "path": "/api/v1/health", "expected_status": 200},
    )

    assert result["result"]["status"] == "succeeded"
    assert result["result"]["status_code"] == 200
    assert result["result"]["assertion_passed"] is True


def test_browser_preview_checks_access_with_fallback_details() -> None:
    db = _memory_session()
    user = _user()
    db.add(user)
    db.commit()

    result = invoke_tool(db, user, "browser.preview", {"url": "/api/v1/health"})

    assert result["result"]["status"] == "succeeded"
    assert result["result"]["accessible"] is True
    assert result["result"]["capability_level"] in {"real", "fallback"}
    assert "playwright" in result["result"]


def test_builtin_tools_sync_to_database_catalog() -> None:
    db = _memory_session()

    sync_builtin_tool_definitions(db)
    tool = db.scalar(select(ToolDefinition).where(ToolDefinition.name == "api.test"))

    assert tool is not None
    assert tool.type == "builtin"
    assert tool.is_builtin is True
    assert tool.builtin_handler == "api.test"
    assert tool.owner_id is None


@pytest.mark.asyncio
async def test_agent_custom_python_tool_is_exposed_and_executed() -> None:
    db = _memory_session()
    user = _user()
    custom_tool = ToolDefinition(
        owner_id=user.id,
        name="agent.custom_echo",
        display_name="Agent Custom Echo",
        description="Echoes text for an agent.",
        type="custom_python",
        status="active",
        input_schema={
            "type": "object",
            "properties": {"input": {"type": "string"}},
            "required": ["input"],
        },
        implementation={
            "language": "python",
            "code": "text = str(arguments.get('input') or '')\nresult = {'echo': text, 'upper': text.upper()}",
        },
    )
    agent = Agent(
        owner_id=user.id,
        name="Custom Tool Agent",
        type="custom",
        config={"tools": [custom_tool.name]},
        capabilities=[],
    )
    conversation = Conversation(
        creator_id=user.id,
        chat_type="single",
        title="Custom Tool Conversation",
    )
    db.add_all([user, custom_tool, agent, conversation])
    db.commit()

    exposed_tools = build_tools_for_agent(db, agent)
    exposed_names = {item["function"]["name"] for item in exposed_tools}

    assert custom_tool.name in exposed_names

    result = await execute_tool_by_name(
        db,
        agent=agent,
        user=user,
        conversation=conversation,
        tool_name=custom_tool.name,
        arguments={"input": "agenthub"},
    )
    invocation = db.scalar(select(ToolInvocation).where(ToolInvocation.tool_name == custom_tool.name))

    assert result["status"] == "succeeded"
    assert result["output"]["result"]["echo"] == "agenthub"
    assert result["output"]["result"]["upper"] == "AGENTHUB"
    assert invocation is not None
    assert invocation.status == "succeeded"

    blocked_agent = Agent(
        owner_id=user.id,
        name="Blocked Custom Tool Agent",
        type="custom",
        config={"tools": [], "capability_permissions_initialized": True},
        capabilities=[],
    )
    db.add(blocked_agent)
    db.commit()

    blocked = await execute_tool_by_name(
        db,
        agent=blocked_agent,
        user=user,
        conversation=conversation,
        tool_name=custom_tool.name,
        arguments={"input": "blocked"},
    )
    invocations = db.scalars(select(ToolInvocation).where(ToolInvocation.tool_name == custom_tool.name)).all()

    assert blocked["type"] == "tool"
    assert blocked["status"] == "failed"
    assert "授权" in blocked["output"]
    assert len(invocations) == 1


def test_cleanup_acceptance_residue_soft_deletes_catalog_noise() -> None:
    db = _memory_session()
    db.add_all(
        [
            McpServer(name="Acceptance HTTP MCP", transport="httpStream", url="http://127.0.0.1", enabled=True),
            ToolDefinition(name="custom_echo_acceptance", display_name="custom_echo_acceptance", type="custom_python"),
            Skill(name="Release Notes Skill", source="ai", status="active"),
            Skill(name="需求分析 Skill", source="system", status="active"),
        ]
    )
    db.commit()

    cleanup_acceptance_residue(db)

    assert db.scalar(select(McpServer).where(McpServer.name == "Acceptance HTTP MCP")).deleted_at is not None
    assert db.scalar(select(ToolDefinition).where(ToolDefinition.name == "custom_echo_acceptance")).status == "deleted"
    assert db.scalar(select(Skill).where(Skill.name == "Release Notes Skill")).status == "deleted"
    assert db.scalar(select(Skill).where(Skill.name == "需求分析 Skill")).status == "active"


def test_mcp_catalog_discovery_and_probe() -> None:
    server = McpServer(
        owner_id="user-1",
        name="Filesystem MCP",
        transport="stdio",
        command="agenthub-mcp-filesystem",
        tool_filter=["file.*"],
        enabled=True,
    )

    tools = discover_server_tools(server)
    probe_server(server)

    assert tools[0]["name"] == "file.*"
    assert tool_allowed(server, "file.read")
    assert server.health_status == "online"


@pytest.mark.asyncio
async def test_mcp_http_probe_discovers_tools_list_and_records_metadata() -> None:
    with _mcp_tools_list_server() as url:
        server = McpServer(
            owner_id="user-1",
            name="HTTP MCP",
            transport="httpStream",
            url=url,
            enabled=True,
        )

        await probe_server_async(server, timeout_ms=1000)

    assert server.health_status == "online"
    assert server.tools == [
        {
            "name": "echo.ping",
            "description": "Echo input",
            "enabled": True,
            "input_schema": {
                "type": "object",
                "properties": {"input": {"type": "string"}},
                "required": ["input"],
            },
        }
    ]
    assert server.extra["last_probe"]["status"] == "ok"
    assert server.extra["last_probe"]["source"] == "http_tools_list"


@pytest.mark.asyncio
async def test_mcp_probe_failure_records_diagnostic_metadata() -> None:
    server = McpServer(
        owner_id="user-1",
        name="Broken MCP",
        transport="httpStream",
        url="http://127.0.0.1:9876",
        enabled=True,
    )

    with patch(
        "app.services.mcp.discovery.list_http_mcp_tools",
        new=AsyncMock(side_effect=ValidationAppError("MCP HTTP tools/list timed out after 5ms")),
    ):
        await probe_server_async(server, timeout_ms=5)

    assert server.health_status == "offline"
    assert server.extra["last_probe"]["status"] == "failed"
    assert server.extra["last_probe"]["source"] == "ValidationAppError"
    assert "timed out" in server.extra["last_probe"]["error"]


def test_mcp_argument_schema_validation() -> None:
    server = McpServer(
        id="server-schema",
        owner_id="user-1",
        name="Schema MCP",
        transport="httpStream",
        url="http://127.0.0.1:8765",
        enabled=True,
        tools=[
            {
                "name": "echo.ping",
                "enabled": True,
                "input_schema": {
                    "type": "object",
                    "properties": {"input": {"type": "string"}},
                    "required": ["input"],
                },
            }
        ],
    )

    validate_mcp_arguments(server, "echo.ping", {"input": "hello"})


@pytest.mark.asyncio
async def test_mcp_allowlist_denial_is_recorded_and_not_executed() -> None:
    db = _memory_session()
    user = _user()
    server = McpServer(
        owner_id=user.id,
        name="Allowlist MCP",
        transport="httpStream",
        url="http://127.0.0.1:9876",
        enabled=True,
        tools=[{"name": "echo.secret", "enabled": False}],
        tool_filter=["echo.*"],
    )
    db.add_all([user, server])
    db.commit()

    with patch("app.services.mcp.invocation.call_http_mcp", new_callable=AsyncMock) as call:
        result = await invoke_mcp_tool_recorded(
            db,
            server=server,
            tool_name_value="echo.secret",
            arguments={"input": "hello"},
            user=user,
        )

    invocation = db.scalar(select(McpToolInvocation).where(McpToolInvocation.server_id == server.id))
    assert result["status"] == "failed"
    assert result["error_code"] == "mcp_tool_not_allowed"
    assert invocation.status == "failed"
    assert invocation.extra["error_code"] == "mcp_tool_not_allowed"
    assert server.health_status == "online"
    call.assert_not_called()


@pytest.mark.asyncio
async def test_mcp_argument_validation_failure_is_recorded() -> None:
    db = _memory_session()
    user = _user()
    server = McpServer(
        owner_id=user.id,
        name="Schema MCP",
        transport="httpStream",
        url="http://127.0.0.1:9876",
        enabled=True,
        tools=[
            {
                "name": "echo.ping",
                "enabled": True,
                "input_schema": {
                    "type": "object",
                    "properties": {"input": {"type": "string"}},
                    "required": ["input"],
                },
            }
        ],
    )
    db.add_all([user, server])
    db.commit()

    with patch("app.services.mcp.invocation.call_http_mcp", new_callable=AsyncMock) as call:
        result = await invoke_mcp_tool_recorded(
            db,
            server=server,
            tool_name_value="echo.ping",
            arguments={},
            user=user,
        )

    assert result["status"] == "failed"
    assert result["error_code"] == "mcp_argument_validation_failed"
    assert db.scalar(select(McpToolInvocation)).result["error_code"] == "mcp_argument_validation_failed"
    assert server.health_status == "online"
    call.assert_not_called()


@pytest.mark.asyncio
async def test_mcp_transport_timeout_is_recorded_with_error_code() -> None:
    db = _memory_session()
    user = _user()
    server = McpServer(
        owner_id=user.id,
        name="Timeout MCP",
        transport="httpStream",
        url="http://127.0.0.1:9876",
        enabled=True,
        tools=[{"name": "echo.ping", "enabled": True}],
        timeout_ms=5,
    )
    db.add_all([user, server])
    db.commit()

    with patch(
        "app.services.mcp.invocation.call_http_mcp",
        new=AsyncMock(side_effect=ValidationAppError("MCP HTTP tool invocation timed out after 5ms")),
    ):
        result = await invoke_mcp_tool_recorded(
            db,
            server=server,
            tool_name_value="echo.ping",
            arguments={"input": "hello"},
            user=user,
            timeout_ms=5,
        )

    invocation = db.scalar(select(McpToolInvocation).where(McpToolInvocation.server_id == server.id))
    assert result["status"] == "failed"
    assert result["error_code"] == "mcp_timeout"
    assert "timed out" in result["error_message"]
    assert invocation.duration_ms >= 0
    assert server.health_status == "offline"


@pytest.mark.asyncio
async def test_mcp_sse_ws_transport_returns_clear_degraded_invocation() -> None:
    db = _memory_session()
    user = _user()
    server = McpServer(
        owner_id=user.id,
        name="Stream MCP",
        transport="sse",
        url="http://127.0.0.1:9876",
        enabled=True,
        tools=[{"name": "echo.stream", "enabled": True}],
    )
    db.add_all([user, server])
    db.commit()

    result = await invoke_mcp_tool_recorded(
        db,
        server=server,
        tool_name_value="echo.stream",
        arguments={"input": "hello"},
        user=user,
    )

    assert result["status"] == "failed"
    assert result["error_code"] == "mcp_transport_unsupported"
    assert "SSE/WebSocket MCP transport" in result["error_message"]
    assert db.scalar(select(McpToolInvocation)).extra["error_code"] == "mcp_transport_unsupported"


def test_import_server_manifest_from_json() -> None:
    manifest = import_server_manifest(
        "json",
        '{"name":"Local MCP","transport":"httpStream","url":"http://127.0.0.1:8765","tools":[{"name":"echo.ping"}]}',
    )

    assert manifest["name"] == "Local MCP"
    assert manifest["tools"][0]["name"] == "echo.ping"


def test_skill_package_parse_and_activation_context(tmp_path: Path) -> None:
    package_dir = tmp_path / "release-skill"
    (package_dir / "scripts").mkdir(parents=True)
    (package_dir / "references").mkdir()
    (package_dir / "SKILL.md").write_text("Use this skill to write release notes.", encoding="utf-8")
    (package_dir / "manifest.json").write_text(
        '{"name":"Release Notes","version":"1.0.1","dependencies":{"tools":["document.review"]}}',
        encoding="utf-8",
    )
    (package_dir / "scripts" / "format.py").write_text("result = arguments", encoding="utf-8")

    manifest = parse_skill_package(package_dir)
    db = _memory_session()
    user = _user()
    skill = Skill(
        id="skill-release",
        owner_id=user.id,
        name="Release Notes",
        description="Write release notes",
        prompt="Use this skill to write release notes.",
        extra={"manifest": manifest},
    )
    agent = Agent(
        owner_id=user.id,
        name="Writer",
        type="custom",
        description="Writer",
        capabilities=[],
        config={"skill_ids": [skill.id]},
    )
    db.add_all([user, skill, agent])
    db.commit()

    context = activated_skill_context(db, agent)

    assert manifest["metadata"]["scripts"] == ["scripts/format.py"]
    assert "Release Notes" in context
    assert f"skill.{skill.id}" in context


def _document_model() -> dict[str, Any]:
    return {
        "kind": "document",
        "title": "结构化项目方案",
        "subtitle": "用于答辩演示的正式方案",
        "template": "proposal",
        "sections": [
            {
                "title": "实施范围",
                "blocks": [
                    {"type": "paragraph", "text": "本方案覆盖需求梳理、开发实施和验收交付。"},
                    {"type": "list", "ordered": True, "items": ["需求确认", "迭代开发", "上线验收"]},
                    {
                        "type": "table",
                        "headers": ["阶段", "交付物"],
                        "rows": [["设计", "方案文档"], ["实现", "可运行平台"]],
                    },
                    {"type": "callout", "title": "注意", "text": "所有产物必须来自真实工具执行。"},
                    {"type": "page_break"},
                    {"type": "heading", "level": 2, "text": "验收标准"},
                    {"type": "paragraph", "text": "支持中文、表格、分页、页眉页脚和页码。"},
                ],
            }
        ],
    }


def test_short_prompt_expands_to_rich_proposal_document() -> None:
    model = normalize_document_model(
        None,
        title="结构化项目方案",
        source_text="结构化项目方案",
        template=None,
    )

    section_titles = [section["title"] for section in model["sections"]]
    first_section_blocks = model["sections"][0]["blocks"]

    assert model["template"] == "proposal"
    assert "项目背景" in section_titles
    assert "实施计划" in section_titles
    assert any(block["type"] == "callout" for block in first_section_blocks)
    assert any(block["type"] == "table" for block in model["sections"][1]["blocks"])


def _memory_session() -> Any:
    engine = create_engine("sqlite:///:memory:", future=True)
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    return session_factory()


def _user() -> User:
    return User(
        id="user-1",
        email="capabilities@example.com",
        username="capabilities",
        password_hash="x",
        display_name="Capabilities",
    )
